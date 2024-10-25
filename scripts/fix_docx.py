import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import argparse
import os

# Set up command-line argument parsing
parser = argparse.ArgumentParser(description="Fix DOCX formatting issues.")
parser.add_argument('--input', type=str, required=True, help='Path to the input DOCX file.')
parser.add_argument('--output', type=str, help='Path to save the modified DOCX file. If not provided, will save as input_filename_output.docx.')
args = parser.parse_args()

# Set output path based on input if not provided
if args.output is None:
    input_path = os.path.abspath(args.input)
    input_dir, input_filename = os.path.split(input_path)
    filename, ext = os.path.splitext(input_filename)
    args.output = os.path.join(input_dir, f"{filename}_output{ext}")

# Set up logging to a file with UTF-8 encoding
log_file = open('conversion_log.txt', 'w', encoding='utf-8')

def log_message(message):
    """Write a message to the log file."""
    log_file.write(message + '\n')
    print(message)  # Optional: print to console as well if desired

def log_exception(exception):
    """Log an exception with its traceback."""
    import traceback
    log_message(f"Exception: {exception}")
    log_message(traceback.format_exc())

# Load the document
try:
    doc = docx.Document(args.input)
    log_message(f"Document '{args.input}' loaded successfully.")
except Exception as e:
    log_exception(e)

bullet_like_pattern = r"^[\u2022\-\*]|^\d+\."  # Pattern to detect bullet point markers
exclude_pattern = r'\.(jpg|png|jpeg|gif)$'  # Exclude common image file extensions

# Function to add missing numbering properties to paragraphs styled as primary or secondary bullets
def add_numbering_properties(paragraph, level):
    try:
        # Create the numbering properties XML elements
        num_pr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), str(level))  # Set level 0 for primary, 1 for secondary bullets
        num_id = OxmlElement('w:numId')
        num_id.set(qn('w:val'), '21' if level == 1 else '19')  # Use appropriate numId for list levels

        # Append numbering properties to the paragraph properties if missing
        pPr = paragraph._element.get_or_add_pPr()
        existing_numPr = pPr.find(qn('w:numPr'))

        if existing_numPr is None:
            pPr.append(num_pr)
            num_pr.append(ilvl)
            num_pr.append(num_id)
            log_message(f"Added numbering properties to bullet point: '{paragraph.text}' at level {level}")
        else:
            # Update existing numbering properties if needed
            existing_ilvl = existing_numPr.find(qn('w:ilvl'))
            if existing_ilvl is not None:
                existing_ilvl.set(qn('w:val'), str(level))
            else:
                existing_numPr.append(ilvl)
            existing_numId = existing_numPr.find(qn('w:numId'))
            if existing_numId is not None:
                existing_numId.set(qn('w:val'), '21' if level == 1 else '19')
            else:
                existing_numPr.append(num_id)
    except Exception as e:
        log_exception(e)

# Function to check if the paragraph is styled as Tiretmoyen using XML
def is_tiretmoyen(paragraph):
    try:
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None and pStyle.get(qn('w:val')) == 'Tiretmoyen':
                return True
    except Exception as e:
        log_exception(e)
    return False

# Function to adjust hyperlink styles within the hyperlink runs only
def adjust_hyperlink_styles(paragraph):
    try:
        # Iterate through XML elements to find hyperlinks
        for hyperlink in paragraph._element.findall(qn('w:hyperlink')):
            # Iterate through runs inside the hyperlink
            for run in hyperlink.findall(qn('w:r')):
                rPr = run.find(qn('w:rPr'))
                if rPr is not None:
                    # Remove unwanted formatting attributes
                    for tag in ['w:b', 'w:sz', 'w:szCs', 'w:i', 'w:strike', 'w:color']:
                        element = rPr.find(qn(tag))
                        if element is not None:
                            rPr.remove(element)
                            log_message(f"Removed '{tag}' formatting from hyperlink text: '{run.text}'")

                    # Ensure the run style is set to 'Hyperlien' and adjust other necessary properties
                    rStyle = rPr.find(qn('w:rStyle'))
                    if rStyle is None:
                        rStyle = OxmlElement('w:rStyle')
                        rPr.append(rStyle)
                    rStyle.set(qn('w:val'), 'Hyperlien')
                    log_message(f"Ensured 'Hyperlien' style for hyperlink text: '{run.text}'")
    except Exception as e:
        log_exception(e)

# Function to clean paragraph properties to resemble the working example
def clean_paragraph_properties(paragraph):
    try:
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            # Remove unnecessary properties that do not exist in the working example
            for tag in ['w:pStyle', 'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:jc']:
                element = pPr.find(qn(tag))
                if element is not None:
                    pPr.remove(element)
                    log_message(f"Removed '{tag}' from paragraph properties to match the working example.")
    except Exception as e:
        log_exception(e)

# Function to adjust paragraph style and convert Tiretmoyen to a secondary bullet point
def adjust_paragraph(paragraph):
    try:
        original_style = paragraph.style.name
        original_text = paragraph.text.strip()  # Trim whitespace around text

        # Exclude paragraphs that look like image filenames or non-bullet text
        if re.search(exclude_pattern, original_text, re.IGNORECASE):
            log_message(f"Skipping adjustment for non-bullet text: '{original_text}'")
            return

        # Check if the paragraph should be styled as a primary or secondary bullet
        is_primary_bullet = original_style == 'Bullet' or re.match(bullet_like_pattern, original_text)
        is_secondary_bullet = original_style.lower() == 'tiretmoyen' or is_tiretmoyen(paragraph)

        # Adjust paragraph style and numbering properties based on type
        if is_primary_bullet:
            log_message(f"Adjusting style for primary bullet point: '{original_text}' to 'List Paragraph'.")
            paragraph.style = 'List Paragraph'
            add_numbering_properties(paragraph, 0)  # Primary bullet level
        elif is_secondary_bullet:
            log_message(f"Converting Tiretmoyen to secondary bullet point: '{original_text}'.")
            # Convert the style to 'Paragraphedeliste' and add numbering properties
            paragraph.style = 'Paragraphedeliste'
            add_numbering_properties(paragraph, 1)  # Secondary bullet level
        else:
            # Clean paragraph properties and adjust hyperlink styles in the paragraph
            clean_paragraph_properties(paragraph)
            adjust_hyperlink_styles(paragraph)
    except Exception as e:
        log_exception(e)

# Loop through each table in the document and adjust only the third column
try:
    for table in doc.tables:
        for row in table.rows:
            # Check if the row has a third cell (index 2)
            if len(row.cells) > 2:
                cell = row.cells[2]

                # Adjust each paragraph within the third column
                for paragraph in cell.paragraphs:
                    adjust_paragraph(paragraph)
except Exception as e:
    log_exception(e)

# Save the modified document
try:
    doc.save(args.output)
    log_message(f"Changes saved to '{args.output}'")
except Exception as e:
    log_exception(e)

# Close the log file
log_file.close()
